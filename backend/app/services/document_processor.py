"""
Document processing service for the Legal Intelligence Platform.

This module handles document text extraction, preprocessing,
and content analysis for various file formats.

Author: Legal Intelligence Platform Team
Version: 1.0.0
"""

import os
import logging
from typing import Optional, Dict, Any
import PyPDF2
from docx import Document as DocxDocument
import re

# Configure logging
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Service class for processing legal documents.
    
    This service handles text extraction from various file formats,
    content preprocessing, and metadata extraction.
    """
    
    def __init__(self):
        """Initialize the document processor."""
        self.supported_formats = {
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.doc': self._extract_doc_text,
            '.txt': self._extract_txt_text
        }
    
    async def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extract text content from a document file.
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            str: Extracted text content, or None if extraction fails
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        # Get file extension
        _, ext = os.path.splitext(file_path.lower())
        
        if ext not in self.supported_formats:
            logger.error(f"Unsupported file format: {ext}")
            return None
        
        try:
            # Extract text using appropriate method
            extractor = self.supported_formats[ext]
            text = await extractor(file_path)
            
            if text:
                # Clean and preprocess text
                text = self._preprocess_text(text)
                logger.info(f"Successfully extracted {len(text)} characters from {file_path}")
                return text
            else:
                logger.warning(f"No text extracted from {file_path}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    async def _extract_pdf_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from PDF file.
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text content
        """
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
            
            return '\n\n'.join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            return None
    
    async def _extract_docx_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from DOCX file.
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n\n'.join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            return None
    
    async def _extract_doc_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from DOC file.
        
        Note: This is a placeholder implementation.
        For production use, consider using python-docx2txt or similar libraries.
        
        Args:
            file_path (str): Path to DOC file
            
        Returns:
            str: Extracted text content
        """
        try:
            # This is a simplified implementation
            # In production, you would use a proper DOC parser
            logger.warning(f"DOC format support is limited for {file_path}")
            return "DOC file content extraction not fully implemented"
            
        except Exception as e:
            logger.error(f"Error reading DOC file {file_path}: {e}")
            return None
    
    async def _extract_txt_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from TXT file.
        
        Args:
            file_path (str): Path to TXT file
            
        Returns:
            str: File content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading TXT file with latin-1 encoding {file_path}: {e}")
                return None
                
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            return None
    
    def _preprocess_text(self, text: str) -> str:
        """
        Preprocess extracted text for analysis.
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            str: Preprocessed text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page numbers and headers/footers (simple heuristic)
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            
            # Skip likely page numbers
            if re.match(r'^\d+$', line):
                continue
            
            # Skip very short lines that might be headers/footers
            if len(line) < 3:
                continue
            
            cleaned_lines.append(line)
        
        # Rejoin lines
        text = '\n'.join(cleaned_lines)
        
        # Remove multiple consecutive newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from document file.
        
        Args:
            file_path (str): Path to document file
            
        Returns:
            dict: Document metadata
        """
        metadata = {
            'file_size': 0,
            'page_count': 0,
            'word_count': 0,
            'character_count': 0,
            'file_type': '',
            'creation_date': None,
            'modification_date': None
        }
        
        try:
            # Basic file information
            stat = os.stat(file_path)
            metadata['file_size'] = stat.st_size
            metadata['modification_date'] = stat.st_mtime
            
            # Get file extension
            _, ext = os.path.splitext(file_path.lower())
            metadata['file_type'] = ext
            
            # Extract text to count words and characters
            text = await self.extract_text(file_path)
            if text:
                metadata['character_count'] = len(text)
                metadata['word_count'] = len(text.split())
            
            # Format-specific metadata
            if ext == '.pdf':
                metadata.update(await self._extract_pdf_metadata(file_path))
            elif ext == '.docx':
                metadata.update(await self._extract_docx_metadata(file_path))
            
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {e}")
        
        return metadata
    
    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF-specific metadata."""
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Extract PDF metadata
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata['title'] = pdf_meta.get('/Title', '')
                    metadata['author'] = pdf_meta.get('/Author', '')
                    metadata['subject'] = pdf_meta.get('/Subject', '')
                    metadata['creator'] = pdf_meta.get('/Creator', '')
                    
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
        
        return metadata
    
    async def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DOCX-specific metadata."""
        metadata = {}
        
        try:
            doc = DocxDocument(file_path)
            
            # Count paragraphs as rough page estimate
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            metadata['page_count'] = max(1, paragraph_count // 25)  # Rough estimate
            
            # Extract document properties
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata['title'] = props.title or ''
                metadata['author'] = props.author or ''
                metadata['subject'] = props.subject or ''
                metadata['creation_date'] = props.created
                
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
        
        return metadata
    
    def get_content_preview(self, text: str, max_length: int = 500) -> str:
        """
        Generate a preview of document content.
        
        Args:
            text (str): Full document text
            max_length (int): Maximum length of preview
            
        Returns:
            str: Content preview
        """
        if not text:
            return ""
        
        if len(text) <= max_length:
            return text
        
        # Try to break at sentence boundary
        preview = text[:max_length]
        last_period = preview.rfind('.')
        last_newline = preview.rfind('\n')
        
        break_point = max(last_period, last_newline)
        
        if break_point > max_length * 0.7:  # If break point is reasonably close to end
            preview = text[:break_point + 1]
        else:
            preview = text[:max_length] + "..."
        
        return preview.strip()
